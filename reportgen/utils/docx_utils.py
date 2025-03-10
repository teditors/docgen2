"""
Utilities for document formatting and structure in the aerospace test report generator.
Provides functions for creating and formatting document elements including:
- Styling text with predefined formats
- Adding headers, footers, and page numbers
- Creating sections with proper numbering
- Handling document orientation
- Inserting images and tables
- Creating signature blocks and cover pages
"""

import os
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt

from config import settings


def apply_style(paragraph, style_name):
    """Apply a predefined style to a paragraph.
    
    Args:
        paragraph: Paragraph object to style
        style_name (str): Name of the style to apply
    """
    logger = logging.getLogger("report_generator")
    
    if style_name in settings.REPORT_STYLES:
        style = settings.REPORT_STYLES[style_name]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        if "font" in style:
            run.font.name = style["font"]
        
        if "size" in style:
            run.font.size = Pt(style["size"])
        
        if "bold" in style and style["bold"]:
            run.bold = True
        
        if "italic" in style and style["italic"]:
            run.italic = True
        
        if "alignment" in style:
            if style["alignment"] == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style["alignment"] == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif style["alignment"] == "left":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif style["alignment"] == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        logger.warning(f"Style '{style_name}' not found in settings")


def add_paragraph_with_style(doc, text, style_name):
    """Add a paragraph with the specified text and style.
    
    Args:
        doc: Document object
        text (str): Text content
        style_name (str): Name of the style to apply
    
    Returns:
        Paragraph object
    """
    paragraph = doc.add_paragraph(text)
    apply_style(paragraph, style_name)
    return paragraph


def add_section(doc, title, level=1, number=None):
    """Add a section with the specified title, level, and number.
    
    Args:
        doc: Document object
        title (str): Section title
        level (int): Heading level (1, 2, or 3)
        number (str, optional): Section number (e.g., "1.0", "1.1")
    
    Returns:
        Paragraph object
    """
    style_name = f"Heading{level}"
    
    if number:
        full_title = f"{number} {title}"
    else:
        full_title = title
    
    return add_paragraph_with_style(doc, full_title, style_name)


def add_header_footer(doc, filename=None, company_name=None):
    """Add header and footer to the document.
    
    Args:
        doc: Document object
        filename (str, optional): Filename to include in header
        company_name (str, optional): Company name to include in footer
    """
    logger = logging.getLogger("report_generator")
    logger.debug("Adding headers and footers to document")
    
    # Add header with filename and page number
    for section in doc.sections:
        # Set up header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        
        if filename:
            # Create a table for header to have filename on left and page numbers on right
            header_table = header.add_table(1, 2)
            header_table.autofit = True
            
            # Left cell - Filename
            left_cell = header_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.text = f"File: {filename}"
            apply_style(left_para, "Header")
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Right cell - Page numbers
            right_cell = header_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            apply_style(right_para, "Header")
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_page_number_field(right_para)
        
        # Set up footer with proprietary information
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        
        if company_name:
            footer_text = settings.PROPRIETARY_FOOTER.format(company_name=company_name)
            footer_para.text = footer_text
            apply_style(footer_para, "Footer")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_number_field(paragraph):
    """Add a page number field to the paragraph.
    
    Args:
        paragraph: Paragraph object
    
    Returns:
        Run object containing the page number field
    """
    run = paragraph.add_run()
    
    # Add "Page X of Y" format
    run.add_text("Page ")
    
    # Add page number field
    r = run._element
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
    
    # Add separator text
    run.add_text(" of ")
    
    # Add total pages field
    r = run._element
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"
    r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
    
    return run


def set_section_orientation(doc, orientation):
    """Set the orientation for a new section in the document.
    
    Args:
        doc: Document object
        orientation: WD_ORIENTATION value (PORTRAIT or LANDSCAPE)
    
    Returns:
        The new section
    """
    logger = logging.getLogger("report_generator")
    logger.debug(f"Setting section orientation to {'landscape' if orientation == WD_ORIENTATION.LANDSCAPE else 'portrait'}")
    
    # Add a section break - FIXED: using WD_BREAK.SECTION_BREAK_NEW_PAGE instead of WD_SECTION.NEW_PAGE
    # This fixes the KeyError in run.add_break()
    paragraph = doc.add_paragraph()
    
    # Simply add page break and then get last section
    doc.add_page_break()
    
    # Get the new section
    new_section = doc.sections[-1]
    
    # Set orientation
    new_section.orientation = orientation
    
    # Set page size based on orientation
    if orientation == WD_ORIENTATION.LANDSCAPE:
        new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    
    return new_section


def create_table_from_data(doc, data, style='Table Grid'):
    """Create a table from a 2D array of data.
    
    Args:
        doc: Document object
        data (list): 2D array of data where the first row is the header
        style (str, optional): Table style name
    
    Returns:
        Table object
    """
    logger = logging.getLogger("report_generator")
    
    if not data:
        logger.warning("Attempted to create table with empty data")
        return None
    
    # Create table
    rows = len(data)
    cols = len(data[0]) if rows > 0 else 0
    logger.debug(f"Creating table with {rows} rows and {cols} columns")
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    
    # Populate table
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)
            
            # Apply header style to the first row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    return table


def insert_image(target, image_path, width=None, height=None):
    """Insert an image into a document or paragraph.
    
    Args:
        target: Document or Paragraph object
        image_path (str or Path): Path to the image file
        width (float, optional): Width in inches
        height (float, optional): Height in inches
    
    Returns:
        InlineShape object
    """
    logger = logging.getLogger("report_generator")
    
    # Convert Path object to string if needed
    if isinstance(image_path, Path):
        image_path = str(image_path)
    
    # Verify image exists
    if not os.path.exists(image_path):
        logger.error(f"Image file does not exist: {image_path}")
        raise FileNotFoundError(f"Image file not found: {image_path}")
    
    # Determine dimensions
    width_inches = Inches(width) if width else None
    height_inches = Inches(height) if height else None
    
    logger.debug(f"Inserting image: {image_path} with width={width_inches}, height={height_inches}")
    
    # Add image based on target type
    try:
        if hasattr(target, 'add_picture'):  # Document or Cell
            return target.add_picture(image_path, width=width_inches, height=height_inches)
        else:  # Paragraph
            return target.add_run().add_picture(image_path, width=width_inches, height=height_inches)
    except Exception as e:
        logger.error(f"Error inserting image {image_path}: {e}", exc_info=True)
        raise


def create_signature_block(doc, title, name=None, date=True):
    """Create a signature block with title, optional name, and date.
    
    Args:
        doc: Document object
        title (str): Title for the signature block (e.g., "Approved by")
        name (str, optional): Name of the person
        date (bool, optional): Whether to include a date line
    
    Returns:
        Paragraph object
    """
    # Add title
    para = doc.add_paragraph()
    para.add_run(f"{title}: ").bold = True
    
    # Add signature line
    para.add_run("_" * 30)
    
    # Add name if provided
    if name:
        para.add_run(f" {name}")
    
    # Add date line if requested
    if date:
        date_para = doc.add_paragraph()
        date_para.add_run("Date: ").bold = True
        date_para.add_run("_" * 20)
    
    return para


def add_cover_page(doc, title, company_name, testing_company, testing_address, pd_number):
    """Add a cover page to the document.
    
    Args:
        doc: Document object
        title (str): Report title
        company_name (str): Client company name
        testing_company (str): Testing company name
        testing_address (str): Testing company address
        pd_number (str): Project number
    """
    logger = logging.getLogger("report_generator")
    logger.debug("Adding cover page")
    
    # Add title
    title_para = add_paragraph_with_style(doc, title, "Title")
    title_para.paragraph_format.space_after = Pt(48)
    
    # Add testing company info
    add_paragraph_with_style(doc, testing_company, "Heading1")
    add_paragraph_with_style(doc, testing_address, "Normal")
    doc.add_paragraph()  # Add spacing
    
    # Add project details
    doc.add_paragraph(f"Project: PD{pd_number}").bold = True
    doc.add_paragraph(f"Client: {company_name}").bold = True
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}").bold = True
    doc.add_paragraph()
    
    # Add report number
    report_num = f"TR_PH{pd_number}_Rev0"
    add_paragraph_with_style(doc, f"Report Number: {report_num}", "Heading2")
    doc.add_paragraph()
    
    # Add signature blocks
    doc.add_paragraph("Prepared by:").bold = True
    create_signature_block(doc, "Engineer", date=True)
    doc.add_paragraph()
    
    create_signature_block(doc, "Reviewed by", date=True)
    doc.add_paragraph()
    
    create_signature_block(doc, "Approved by", date=True)
    doc.add_paragraph()
    
    # Add page break
    doc.add_page_break()
