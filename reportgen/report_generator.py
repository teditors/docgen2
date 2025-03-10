"""
Core report generation functionality for aerospace hardware test reports.
"""

import os
from datetime import datetime
from pathlib import Path
import tempfile

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from config import settings
from utils.docx_utils import (
    add_cover_page,
    add_header_footer,
    add_paragraph_with_style,
    add_section,
    apply_style,
    create_signature_block,
    create_table_from_data,
    insert_image,
    set_section_orientation,
)
from utils.file_utils import (
    convert_pdf_to_image,
    extract_test_data,
    find_files,
    get_po_details,
    parse_test_logs,
)


class ReportGenerator:
    """Handles the generation of test reports based on project data."""

    def __init__(self, project_dir, pd_number, company_name, test_type, output_dir, logger):
        """Initialize the report generator with project details.
        
        Args:
            project_dir (Path): Path to the project directory
            pd_number (str): Project number
            company_name (str): Client company name
            test_type (str): Type of test report to generate
            output_dir (Path): Directory to save the generated report
            logger: Logger instance
        """
        self.project_dir = project_dir
        self.pd_number = pd_number
        self.company_name = company_name
        self.test_type = test_type
        self.output_dir = output_dir
        self.logger = logger
        
        # Create a persistent temp directory for image conversions
        self.temp_dir = Path(tempfile.gettempdir()) / f"report_gen_{pd_number}"
        self.temp_dir.mkdir(exist_ok=True)
        self.logger.info(f"Using temp directory: {self.temp_dir}")
        
        # Get test folder paths
        if test_type == "all":
            self.test_folders = self._get_all_test_folders()
        else:
            test_path = self.project_dir / "testing" / test_type
            self.test_folders = list(test_path.glob("PHB*"))
        
        self.logger.info(f"Found {len(self.test_folders)} test folders for {test_type}")
        for folder in self.test_folders:
            self.logger.debug(f"  - {folder}")
        
        # Generate output filename
        self.output_filename = f"TR_PH{pd_number}_Rev0_{company_name}_{test_type}.docx"
        self.output_path = output_dir / self.output_filename
    
    def _get_all_test_folders(self):
        """Get all test folders across all test types."""
        all_folders = []
        test_types = ["Dynamics", "EMIEMC", "Environmental"]
        
        for test_type in test_types:
            test_path = self.project_dir / "testing" / test_type
            if test_path.exists():
                folders = list(test_path.glob("PHB*"))
                all_folders.extend(folders)
                self.logger.debug(f"Found {len(folders)} {test_type} test folders")
        
        return all_folders
    
    def generate_report(self):
        """Generate the complete test report.
        
        Returns:
            Path: Path to the generated report file
        """
        self.logger.info("Creating new document")
        doc = Document()
        
        # Add cover page
        self._add_cover_page(doc)
        
        # Add administrative information section
        self._add_admin_section(doc)
        
        # Add test sections
        self._add_test_sections(doc)
        
        # Add ending page
        self._add_end_page(doc)
        
        # Save the document
        self.logger.info(f"Saving document to {self.output_path}")
        doc.save(self.output_path)
        return self.output_path
    
    def _add_cover_page(self, doc):
        """Add the cover page to the document."""
        self.logger.info("Adding cover page")
        
        # Get the descriptive title based on test type
        if self.test_type == "all":
            test_title = "Comprehensive Test Report"
        else:
            test_title = f"{self.test_type} Test Report"
        
        add_cover_page(
            doc,
            title=test_title,
            company_name=self.company_name,
            testing_company=settings.COMPANY_NAME,
            testing_address=settings.COMPANY_ADDRESS,
            pd_number=self.pd_number,
        )
    
    def _add_admin_section(self, doc):
        """Add the administrative information section."""
        self.logger.info("Adding administrative information section")
        
        # Start section 1.0
        add_section(doc, "Administrative Information", level=1, number="1.0")
        
        # Section 1.1 - Units Under Test Table
        self._add_units_under_test_table(doc)
        
        # Section 1.2 - Tests Performed Table
        self._add_tests_performed_table(doc)
        
        # Section 1.3 - Details
        self._add_admin_details(doc)
    
    def _add_units_under_test_table(self, doc):
        """Add the Units Under Test table (Section 1.1)."""
        self.logger.info("Adding units under test table")
        add_section(doc, "Units Under Test", level=2, number="1.1")
        
        # Extract unit information from PO
        try:
            po_files = find_files(self.project_dir / "admin" / "PO", "PO*")
            if po_files:
                units_data = get_po_details(po_files[0])
                # Placeholder for actual PO parsing logic
                # For now, use stub data
                units_data = [
                    ["Item", "Part Number", "Serial Number", "Description", "Quantity"],
                    ["1", "XYZ-123", "SN001", "Controller Assembly", "1"],
                    ["2", "XYZ-456", "SN002", "Sensor Package", "2"],
                ]
                
                create_table_from_data(doc, units_data)
            else:
                self.logger.warning("No PO information found")
                add_paragraph_with_style(doc, "No PO information found.", "Normal")
        except Exception as e:
            self.logger.error(f"Error extracting units from PO: {e}", exc_info=True)
            add_paragraph_with_style(doc, "Error extracting unit information.", "Normal")
    
    def _add_tests_performed_table(self, doc):
        """Add the Tests Performed table (Section 1.2)."""
        self.logger.info("Adding tests performed table")
        add_section(doc, "Tests Performed", level=2, number="1.2")
        
        # Extract test information from test logs
        tests_data = [["Test", "Date Performed", "Specification", "Result", "NODs"]]
        
        try:
            for test_folder in self.test_folders:
                test_logs = find_files(test_folder / "worksheets", "TestLog_*.pdf")
                for log in test_logs:
                    # Parse test log information
                    # This is a stub - actual parsing would be implemented in parse_test_logs
                    test_details = parse_test_logs(log)
                    
                    # Check if NODs exist for this test
                    nods_exist = any(Path(test_folder / "NODs").glob("NOD_*.pdf"))
                    nods_text = "Yes" if nods_exist else "No"
                    
                    # Add test information to the table
                    tests_data.append([
                        test_details.get("test_name", "Unknown Test"),
                        test_details.get("date", "Unknown"),
                        test_details.get("specification", "Not specified"),
                        test_details.get("result", "UNKNOWN"),
                        nods_text
                    ])
            
            if len(tests_data) > 1:
                create_table_from_data(doc, tests_data)
            else:
                self.logger.warning("No test information found")
                add_paragraph_with_style(doc, "No test information found.", "Normal")
        except Exception as e:
            self.logger.error(f"Error creating tests performed table: {e}", exc_info=True)
            add_paragraph_with_style(doc, "Error extracting test information.", "Normal")
    
    def _add_admin_details(self, doc):
        """Add administrative details (Section 1.3)."""
        self.logger.info("Adding administrative details")
        add_section(doc, "Details", level=2, number="1.3")
        
        # Add boilerplate text
        boilerplate = (
            "The test results contained in this report are presented as measured during the "
            "specified test procedures. The interpretation of these results and their "
            "applicability to the final product performance is the responsibility of the customer. "
            f"{settings.COMPANY_NAME} makes no claims regarding the fitness for purpose of the "
            "tested items beyond the specific test conditions described herein."
        )
        
        add_paragraph_with_style(doc, boilerplate, "Normal")
    
    def _add_test_sections(self, doc):
        """Add sections for each test."""
        self.logger.info("Adding test sections")
        
        # Start with section number 2
        section_number = 2
        
        for test_folder in self.test_folders:
            try:
                test_name = self._get_test_name(test_folder)
                self.logger.info(f"Processing test section: {test_name}")
                
                # Add main section for this test
                add_section(doc, test_name, level=1, number=f"{section_number}.0")
                
                # Add subsections
                self._add_test_procedure(doc, test_folder, f"{section_number}.1")
                self._add_test_result(doc, test_folder, f"{section_number}.2")
                self._add_test_datasheets(doc, test_folder, f"{section_number}.3")
                self._add_test_nods(doc, test_folder, f"{section_number}.4")
                self._add_test_photographs(doc, test_folder, f"{section_number}.5")
                self._add_test_plots(doc, test_folder, f"{section_number}.6")
                
                section_number += 1
            except Exception as e:
                self.logger.error(f"Error processing test folder {test_folder}: {e}", exc_info=True)
                add_paragraph_with_style(doc, f"Error processing test: {test_folder.name}", "Normal")
    
    def _get_test_name(self, test_folder):
        """Extract a descriptive name for the test from the folder or files."""
        # Try to get test name from worksheets
        test_logs = list(Path(test_folder / "worksheets").glob("TestLog_*.pdf"))
        if test_logs:
            # Extract component and test procedure from filename
            # Format: TestLog_Component_TestProc_YYYYMMDD.pdf
            filename = test_logs[0].stem
            parts = filename.split('_')
            if len(parts) >= 3:
                return f"{parts[1]} {parts[2]} Test"
        
        # Fallback to folder name
        return f"Test {test_folder.name}"
    
    def _add_test_procedure(self, doc, test_folder, section_number):
        """Add the test procedure subsection."""
        self.logger.debug(f"Adding test procedure for {test_folder.name}")
        add_section(doc, "Procedure", level=2, number=section_number)
        
        # Try to extract procedure from test logs
        test_logs = list(Path(test_folder / "worksheets").glob("TestLog_*.pdf"))
        if test_logs:
            # Placeholder for actual test log parsing
            add_paragraph_with_style(
                doc,
                "This test was conducted according to the procedure specified in the test log.",
                "Normal",
            )
        else:
            self.logger.warning(f"No test logs found in {test_folder / 'worksheets'}")
            add_paragraph_with_style(doc, "No test procedure information available.", "Normal")
    
    def _add_test_result(self, doc, test_folder, section_number):
        """Add the test result subsection."""
        self.logger.debug(f"Adding test result for {test_folder.name}")
        add_section(doc, "Result", level=2, number=section_number)
        
        # Extract result from test logs
        test_logs = list(Path(test_folder / "worksheets").glob("TestLog_*.pdf"))
        if test_logs:
            # Placeholder for actual test log parsing
            test_result = "PASS"  # Default placeholder
            add_paragraph_with_style(doc, f"Test Result: {test_result}", "Normal")
        else:
            self.logger.warning(f"No test logs found in {test_folder / 'worksheets'}")
            add_paragraph_with_style(doc, "No test result information available.", "Normal")
    
    def _add_test_datasheets(self, doc, test_folder, section_number):
        """Add the datasheets subsection."""
        worksheet_path = test_folder / "worksheets"
        if not worksheet_path.exists() or not any(worksheet_path.glob("TestLog_*.pdf")):
            self.logger.debug(f"No worksheets found in {worksheet_path}, skipping datasheets section")
            return  # Skip this section if no worksheets
        
        self.logger.debug(f"Adding datasheets for {test_folder.name}")
        add_section(doc, "Datasheets", level=2, number=section_number)
        
        # Convert PDFs to images and embed
        worksheet_files = list(worksheet_path.glob("TestLog_*.pdf"))
        for worksheet in worksheet_files:
            try:
                self.logger.debug(f"Processing worksheet: {worksheet}")
                
                # Convert PDF to image for embedding - use the class temp directory
                image_path = convert_pdf_to_image(worksheet, output_dir=self.temp_dir / "datasheets")
                
                # Add caption
                add_paragraph_with_style(doc, f"Datasheet: {worksheet.stem}", "Caption")
                
                # Insert image
                insert_image(doc, image_path, width=6.5)
                
                # Add a small space after the image
                doc.add_paragraph()
            except Exception as e:
                self.logger.error(f"Error embedding worksheet {worksheet}: {e}", exc_info=True)
                add_paragraph_with_style(doc, f"Error embedding worksheet: {worksheet.name}", "Normal")
    
    def _add_test_nods(self, doc, test_folder, section_number):
        """Add the NODs (Notice of Deviation) subsection."""
        nods_path = test_folder / "NODs"
        if not nods_path.exists() or not any(nods_path.glob("NOD_*.pdf")):
            self.logger.debug(f"No NODs found in {nods_path}, skipping NODs section")
            return  # Skip this section if no NODs
        
        self.logger.debug(f"Adding NODs for {test_folder.name}")
        add_section(doc, "Notices of Deviation", level=2, number=section_number)
        
        # Convert PDFs to images and embed
        nod_files = list(nods_path.glob("NOD_*.pdf"))
        for nod in nod_files:
            try:
                self.logger.debug(f"Processing NOD: {nod}")
                
                # Convert PDF to image for embedding - use the class temp directory
                image_path = convert_pdf_to_image(nod, output_dir=self.temp_dir / "nods")
                
                # Add caption - extract date from filename (NOD_mm.dd.yyyy.pdf)
                nod_date = nod.stem.replace("NOD_", "")
                add_paragraph_with_style(doc, f"Notice of Deviation: {nod_date}", "Caption")
                
                # Insert image
                insert_image(doc, image_path, width=6.5)
                
                # Add a small space after the image
                doc.add_paragraph()
            except Exception as e:
                self.logger.error(f"Error embedding NOD {nod}: {e}", exc_info=True)
                add_paragraph_with_style(doc, f"Error embedding NOD: {nod.name}", "Normal")
    
    def _add_test_photographs(self, doc, test_folder, section_number):
        """Add the photographs subsection."""
        photos_path = test_folder / "photographs"
        
        # Check if the photographs directory exists
        if not photos_path.exists():
            self.logger.debug(f"Photographs directory does not exist: {photos_path}")
            return  # Skip this section
        
        # Try all possible patterns to find photos
        photo_files = []
        
        # First, try the most specific patterns
        patterns = [
            "photo_*.jpg",
            "photo_*.jpeg",
            "*.jpg",
            "*.jpeg",
        ]
        
        for pattern in patterns:
            found_files = list(photos_path.glob(pattern))
            if found_files:
                photo_files.extend(found_files)
                self.logger.debug(f"Found {len(found_files)} photos with pattern '{pattern}' in {photos_path}")
        
        if not photo_files:
            self.logger.debug(f"No photographs found in {photos_path} after trying all patterns")
            return  # Skip this section if no photographs
        
        # Log the found photos for debugging
        self.logger.info(f"Adding photographs section with {len(photo_files)} photos")
        for photo in photo_files[:5]:  # Log first 5 photos
            self.logger.debug(f"  - {photo.name}")
        if len(photo_files) > 5:
            self.logger.debug(f"  - ... and {len(photo_files) - 5} more")
        
        # Add the section header
        add_section(doc, "Photographs", level=2, number=section_number)
        
        # Add photographs, 2 per page
        # Process photos in pairs
        for i in range(0, len(photo_files), 2):
            table = doc.add_table(rows=2, cols=1)
            table.style = 'Table Grid'
            
            # Add first photo
            cell = table.cell(0, 0)
            try:
                photo = photo_files[i]
                self.logger.debug(f"Processing photo: {photo}")
                
                # Verify the photo exists
                if not photo.exists():
                    self.logger.error(f"Photo file does not exist: {photo}")
                    cell.text = f"Error: Photo file not found: {photo.name}"
                    continue
                
                # Extract photo description from filename (photo_###_Component.jpg)
                filename = photo.stem
                if filename.startswith("photo_"):
                    parts = filename.split("_", 2)  # Split at most 2 times
                    if len(parts) >= 3:
                        photo_desc = parts[2]  # Get the Component part
                    elif len(parts) == 2:
                        photo_desc = parts[1]  # Get the number part as fallback
                    else:
                        photo_desc = filename  # Use full filename as last resort
                else:
                    # If filename doesn't follow the expected pattern
                    photo_desc = filename
                
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Insert image
                insert_image(cell_para, photo, width=3.5)
                
                # Add caption
                caption_para = cell.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_style(caption_para, "Caption")
                caption_para.add_run(f"Photo: {photo_desc}")
            except Exception as e:
                self.logger.error(f"Error embedding photo {photo_files[i]}: {e}", exc_info=True)
                cell.text = f"Error embedding photo: {photo_files[i].name}"
            
            # Add second photo if available
            if i + 1 < len(photo_files):
                cell = table.cell(1, 0)
                try:
                    photo = photo_files[i+1]
                    self.logger.debug(f"Processing photo: {photo}")
                    
                    # Verify the photo exists
                    if not photo.exists():
                        self.logger.error(f"Photo file does not exist: {photo}")
                        cell.text = f"Error: Photo file not found: {photo.name}"
                        continue
                    
                    # Extract photo description from filename
                    filename = photo.stem
                    if filename.startswith("photo_"):
                        parts = filename.split("_", 2)  # Split at most 2 times
                        if len(parts) >= 3:
                            photo_desc = parts[2]  # Get the Component part
                        elif len(parts) == 2:
                            photo_desc = parts[1]  # Get the number part as fallback
                        else:
                            photo_desc = filename  # Use full filename as last resort
                    else:
                        # If filename doesn't follow the expected pattern
                        photo_desc = filename
                    
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Insert image
                    insert_image(cell_para, photo, width=3.5)
                    
                    # Add caption
                    caption_para = cell.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    apply_style(caption_para, "Caption")
                    caption_para.add_run(f"Photo: {photo_desc}")
                except Exception as e:
                    self.logger.error(f"Error embedding photo {photo_files[i+1]}: {e}", exc_info=True)
                    cell.text = f"Error embedding photo: {photo_files[i+1].name}"
            
            # Add page break after each pair unless it's the last pair
            if i + 2 < len(photo_files):
                doc.add_page_break()
    
    def _add_test_plots(self, doc, test_folder, section_number):
        """Add the plots subsection in landscape orientation."""
        data_path = test_folder / "data"
        if not data_path.exists():
            self.logger.debug(f"Data directory {data_path} doesn't exist, skipping plots section")
            return  # Skip this section if no data directory
            
        # Check for plot files with multiple patterns
        plot_files = []
        patterns = [
            "*_Description.jpg",
            "*_Description.jpeg",
            "*.jpg",
            "*.jpeg"
        ]
        
        for pattern in patterns:
            found_files = list(data_path.glob(pattern))
            if found_files:
                plot_files.extend(found_files)
                self.logger.debug(f"Found {len(found_files)} plots with pattern '{pattern}' in {data_path}")
        
        if not plot_files:
            self.logger.debug(f"No plot files found in {data_path} after trying all patterns")
            return  # Skip this section if no plot files
        
        # Log the found plots for debugging
        self.logger.info(f"Adding plots section with {len(plot_files)} plots")
        for plot in plot_files[:5]:  # Log first 5 plots
            self.logger.debug(f"  - {plot.name}")
        if len(plot_files) > 5:
            self.logger.debug(f"  - ... and {len(plot_files) - 5} more")
        
        # Add the section header
        add_section(doc, "Plots", level=2, number=section_number)
        
        # Add a page break when done with plots
        doc.add_page_break()
        
        # Reset to portrait orientation for next sections
        current_section = doc.sections[-1]
        current_section.orientation = WD_ORIENTATION.PORTRAIT
        # Swap width and height back
        current_section.page_width, current_section.page_height = current_section.page_height, current_section.page_width
        
        self.logger.debug(f"Switched back to portrait orientation after plots") #break before switching to landscape
        doc.add_page_break()
        
        # Get the current section
        current_section = doc.sections[-1]
        
        # Set orientation directly without creating a new section
        current_section.orientation = WD_ORIENTATION.LANDSCAPE
        
        # Swap width and height
        current_section.page_width, current_section.page_height = current_section.page_height, current_section.page_width
        
        self.logger.debug(f"Switched to landscape orientation for plots")
        
        # Add plots, one per page
        for i, plot in enumerate(plot_files):
            try:
                self.logger.debug(f"Processing plot: {plot}")
                
                # Verify the plot exists
                if not plot.exists():
                    self.logger.error(f"Plot file does not exist: {plot}")
                    add_paragraph_with_style(doc, f"Error: Plot file not found: {plot.name}", "Normal")
                    continue
                
                # Extract plot description from filename
                filename = plot.stem
                if "_Description" in filename:
                    # Format: ###_Description.jpg
                    parts = filename.split("_Description")
                    plot_desc = f"Plot {parts[0]}"
                else:
                    # Try to extract a meaningful description
                    parts = filename.split("_")
                    if len(parts) > 1:
                        plot_desc = " ".join(parts)
                    else:
                        plot_desc = filename
                
                # Add caption
                add_paragraph_with_style(doc, f"Plot: {plot_desc}", "Caption")
                
                # Insert image - adjust width to fit landscape orientation
                insert_image(doc, plot, width=9.0)  # Wider in landscape mode
                
                # Add page break after each plot unless it's the last one
                if i < len(plot_files) - 1:
                    doc.add_page_break()
            except Exception as e:
                self.logger.error(f"Error embedding plot {plot}: {e}", exc_info=True)
                add_paragraph_with_style(doc, f"Error embedding plot: {plot.name}", "Normal")
        
        # Add a page
    
    def _add_end_page(self, doc):
        """Add the end page to the document."""
        self.logger.info("Adding end page")
        
        doc.add_page_break()
        
        # Add END OF REPORT text
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Position text in the middle of the page
        for _ in range(10):
            paragraph.add_run().add_break()
        
        # Add the text in bold and all caps
        run = paragraph.add_run("END OF REPORT")
        run.bold = True
        run.font.size = Pt(16)